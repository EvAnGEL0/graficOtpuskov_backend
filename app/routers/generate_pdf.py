from fastapi import FastAPI, HTTPException, APIRouter, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.config.database import get_db
from app import models
from sqlalchemy.orm import selectinload
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH  # ✅ Добавьте это
import asyncio

router = APIRouter(
    prefix="/generate_pdf",
    tags=["generate_pdf"]
)

@router.post("/generate-vacation-schedule-docx/")
async def generate_vacation_schedule_docx(department_id: int, year: int, db: AsyncSession = Depends(get_db)):
    try:
        # Получаем имя отдела
        dept_result = await db.execute(
            select(models.Department_s).where(models.Department_s.id == department_id)
        )
        department = dept_result.scalar_one_or_none()

        if not department:
            raise HTTPException(status_code=404, detail="Отдел не найден")

        # Выполняем запрос к базе
        result = await db.execute(
            select(models.VacationSchedule)
            .options(
                selectinload(models.VacationSchedule.staff)
                .selectinload(models.Staff.position)
            )
            .join(models.Staff)
            .where(
                models.Staff.department_id == department_id,
                models.VacationSchedule.start_date >= f"{year}-01-01",
                models.VacationSchedule.end_date <= f"{year}-12-31"
            )
        )
        vacations_db = result.scalars().all()

        if not vacations_db:
            raise HTTPException(status_code=404, detail="Нет данных об отпусках для выбранного отдела и года")

        # Создаём документ
        doc = Document()

        # Заголовок "График отпусков за {year} год" по центру
        heading1 = doc.add_paragraph()
        heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER  # ✅ Центрируем
        heading1.add_run(f'График отпусков за {year} год').bold = True

        # Заголовок "отдел: {department.name}" по центру
        heading2 = doc.add_paragraph()
        heading2.alignment = WD_ALIGN_PARAGRAPH.CENTER  # ✅ Центрируем
        heading2.add_run(f'отдел: {department.name}').bold = True

        # Добавляем таблицу
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Должность'
        hdr_cells[1].text = 'Фамилия'
        hdr_cells[2].text = 'Имя'
        hdr_cells[3].text = 'Отчество'
        hdr_cells[4].text = 'Кол-во дней'
        hdr_cells[5].text = 'Период отпуска'

        # Центрируем текст в заголовках таблицы (опционально)
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for vac in vacations_db:
            row_cells = table.add_row().cells
            row_cells[0].text = vac.staff.position.name if vac.staff.position else "Не указана"
            row_cells[1].text = vac.staff.last_name
            row_cells[2].text = vac.staff.first_name
            row_cells[3].text = vac.staff.middle_name or ""
            row_cells[4].text = str(vac.main_vacation_days)
            row_cells[5].text = f"{vac.start_date.strftime('%d.%m.%Y')} - {vac.end_date.strftime('%d.%m.%Y')}"

        # Подпись внизу по центру
        doc.add_paragraph()  # пустая строка перед подписью
        signature = doc.add_paragraph()
        signature.add_run('Начальник отдела: _____________________________')

        # Сохраняем в BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Возвращаем файл
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=grafik_otpuska_{year}_{department.name}.docx"
            }
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка генерации DOCX: {str(e)}")
    
    
    
    

@router.post("/generate-all-departments-schedule-docx/")
async def generate_all_departments_schedule_docx(year: int, db: AsyncSession = Depends(get_db)):
    try:
        # Получаем все отделы
        dept_result = await db.execute(
            select(models.Department_s)
        )
        departments = dept_result.scalars().all()

        if not departments:
            raise HTTPException(status_code=404, detail="Нет отделов в системе")

        # Создаём документ
        doc = Document()
        doc.add_heading(f'График отпусков за {year} год (все отделы)', 0)

        for dept in departments:
            # Заголовок отдела
            dept_heading = doc.add_paragraph()
            dept_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dept_heading.add_run(f'Отдел: {dept.name}').bold = True

            # Запрос отпусков для отдела
            result = await db.execute(
                select(models.VacationSchedule)
                .options(
                    selectinload(models.VacationSchedule.staff)
                    .selectinload(models.Staff.position)
                )
                .join(models.Staff)
                .where(
                    models.Staff.department_id == dept.id,
                    models.VacationSchedule.start_date >= f"{year}-01-01",
                    models.VacationSchedule.end_date <= f"{year}-12-31"
                )
            )
            vacations_db = result.scalars().all()

            if not vacations_db:
                # Если нет отпусков — пропускаем
                doc.add_paragraph('Нет данных об отпусках.')
            else:
                # Добавляем таблицу
                table = doc.add_table(rows=1, cols=6)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Должность'
                hdr_cells[1].text = 'Фамилия'
                hdr_cells[2].text = 'Имя'
                hdr_cells[3].text = 'Отчество'
                hdr_cells[4].text = 'Кол-во дней'
                hdr_cells[5].text = 'Период отпуска'

                # Центрируем заголовки таблицы (опционально)
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for vac in vacations_db:
                    row_cells = table.add_row().cells
                    row_cells[0].text = vac.staff.position.name if vac.staff.position else "Не указана"
                    row_cells[1].text = vac.staff.last_name
                    row_cells[2].text = vac.staff.first_name
                    row_cells[3].text = vac.staff.middle_name or ""
                    row_cells[4].text = str(vac.main_vacation_days)
                    row_cells[5].text = f"{vac.start_date.strftime('%d.%m.%Y')} - {vac.end_date.strftime('%d.%m.%Y')}"

            # Разделитель между отделами
            doc.add_page_break()  # или doc.add_paragraph() для пустой строки

        # Подпись внизу (опционально — можно убрать или оставить на последней странице)
        # doc.add_paragraph()
        # signature = doc.add_paragraph()
        # signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # signature.add_run('Начальник отдела: _____________________________')

        # Сохраняем в BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Возвращаем файл
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=grafik_otpuska_all_departments_{year}.docx"
            }
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка генерации DOCX: {str(e)}")